"""
Document redlining module for Blackletter Systems.

This module provides functionality for creating redlined documents:
- Create redline DOCX files from original text
- Highlight changes and issues
- Add comments and suggestions

Usage:
    from app.core.redact import create_redlined_document
    
    # Create a redlined document
    docx_bytes = create_redlined_document(
        original_text="This is the original contract text...",
        issues=[
            {
                "text": "This clause is problematic",
                "start": 10,
                "end": 50,
                "severity": "high",
                "comment": "Consider revising this clause"
            }
        ]
    )
"""

import io
from typing import Dict, List, Optional, Tuple, Union, Any
import logging
from enum import Enum

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_COLOR_INDEX, WD_PARAGRAPH_ALIGNMENT

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class IssueSeverity(str, Enum):
    """Issue severity levels"""
    LOW = "low"
    MEDIUM = "medium"
    HIGH = "high"
    CRITICAL = "critical"

class Issue:
    """Represents an issue in the document"""
    def __init__(
        self,
        text: str,
        start: int,
        end: int,
        severity: IssueSeverity = IssueSeverity.MEDIUM,
        comment: Optional[str] = None,
        suggestion: Optional[str] = None
    ):
        self.text = text
        self.start = start
        self.end = end
        self.severity = severity
        self.comment = comment
        self.suggestion = suggestion

def create_redlined_document(
    original_text: str,
    issues: List[Dict[str, Any]],
    title: Optional[str] = None
) -> bytes:
    """
    Create a redlined document with highlighted issues.
    
    Args:
        original_text: The original document text
        issues: List of issues with keys:
               text, start, end, severity, comment, suggestion
        title: Optional document title
        
    Returns:
        bytes: The redlined document as bytes
    """
    # Create a new document
    doc = Document()
    
    # Add title if provided
    if title:
        doc.add_heading(title, level=1)
    
    # Add summary section
    summary_heading = doc.add_heading("Summary of Issues", level=2)
    
    # Group issues by severity
    issues_by_severity = {
        IssueSeverity.CRITICAL: [],
        IssueSeverity.HIGH: [],
        IssueSeverity.MEDIUM: [],
        IssueSeverity.LOW: []
    }
    
    for issue_data in issues:
        severity = issue_data.get("severity", "medium").lower()
        if severity not in [s.value for s in IssueSeverity]:
            severity = "medium"
        
        issue = Issue(
            text=issue_data["text"],
            start=issue_data["start"],
            end=issue_data["end"],
            severity=IssueSeverity(severity),
            comment=issue_data.get("comment"),
            suggestion=issue_data.get("suggestion")
        )
        
        issues_by_severity[IssueSeverity(severity)].append(issue)
    
    # Add summary table
    summary_table = doc.add_table(rows=1, cols=3)
    summary_table.style = "Table Grid"
    
    # Add header row
    header_cells = summary_table.rows[0].cells
    header_cells[0].text = "Severity"
    header_cells[1].text = "Count"
    header_cells[2].text = "Description"
    
    # Add data rows
    for severity in [IssueSeverity.CRITICAL, IssueSeverity.HIGH, IssueSeverity.MEDIUM, IssueSeverity.LOW]:
        issues_count = len(issues_by_severity[severity])
        if issues_count > 0:
            row_cells = summary_table.add_row().cells
            row_cells[0].text = severity.value.capitalize()
            row_cells[1].text = str(issues_count)
            
            if severity == IssueSeverity.CRITICAL:
                row_cells[2].text = "Critical issues requiring immediate attention"
            elif severity == IssueSeverity.HIGH:
                row_cells[2].text = "Significant issues that should be addressed"
            elif severity == IssueSeverity.MEDIUM:
                row_cells[2].text = "Moderate issues to consider"
            else:
                row_cells[2].text = "Minor issues for awareness"
    
    # Add space after summary
    doc.add_paragraph()
    
    # Add document content with redlines
    doc.add_heading("Document with Issues Highlighted", level=2)
    
    # Sort issues by start position (in reverse to avoid index shifting)
    sorted_issues = sorted(
        [Issue(**issue) for issue in issues],
        key=lambda x: x.start,
        reverse=True
    )
    
    # Insert the text with highlights
    text = original_text
    
    # Add the text as paragraphs
    paragraphs = text.split("\n")
    
    # Track the current position in the original text
    current_pos = 0
    
    for para_text in paragraphs:
        if not para_text.strip():
            doc.add_paragraph()
            current_pos += 1  # Account for the newline
            continue
        
        paragraph = doc.add_paragraph()
        
        # Find issues that overlap with this paragraph
        para_start = current_pos
        para_end = para_start + len(para_text)
        
        # Get issues that overlap with this paragraph
        para_issues = []
        for issue in sorted_issues:
            if (issue.start < para_end and issue.end > para_start):
                # Adjust issue boundaries to be relative to the paragraph
                relative_start = max(0, issue.start - para_start)
                relative_end = min(len(para_text), issue.end - para_start)
                
                para_issues.append({
                    "start": relative_start,
                    "end": relative_end,
                    "severity": issue.severity,
                    "comment": issue.comment,
                    "suggestion": issue.suggestion
                })
        
        # Sort paragraph issues by start position (in reverse)
        para_issues.sort(key=lambda x: x["start"], reverse=True)
        
        # Apply highlighting to the paragraph text
        remaining_text = para_text
        
        if not para_issues:
            # No issues in this paragraph, add the text as-is
            paragraph.add_run(remaining_text)
        else:
            # Apply highlighting for each issue
            for issue in para_issues:
                start = issue["start"]
                end = issue["end"]
                
                # Add text after the issue
                if end < len(remaining_text):
                    paragraph.add_run(remaining_text[end:])
                
                # Add the highlighted issue text
                issue_text = remaining_text[start:end]
                run = paragraph.add_run(issue_text)
                
                # Apply formatting based on severity
                if issue["severity"] == IssueSeverity.CRITICAL:
                    run.font.highlight_color = WD_COLOR_INDEX.RED
                    run.bold = True
                elif issue["severity"] == IssueSeverity.HIGH:
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    run.bold = True
                elif issue["severity"] == IssueSeverity.MEDIUM:
                    run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
                else:
                    run.font.highlight_color = WD_COLOR_INDEX.TURQUOISE
                
                # Add text before the issue
                if start > 0:
                    paragraph.add_run(remaining_text[:start])
                
                # Update remaining text for next iteration
                remaining_text = remaining_text[:start]
        
        # Update current position
        current_pos = para_end + 1  # +1 for the newline
    
    # Add detailed issues section
    doc.add_page_break()
    doc.add_heading("Detailed Issues", level=2)
    
    # Sort issues by severity and then by position
    all_issues = []
    for severity in [IssueSeverity.CRITICAL, IssueSeverity.HIGH, IssueSeverity.MEDIUM, IssueSeverity.LOW]:
        all_issues.extend(sorted(
            issues_by_severity[severity],
            key=lambda x: x.start
        ))
    
    # Add each issue
    for i, issue in enumerate(all_issues, 1):
        # Add issue heading
        issue_heading = doc.add_heading(f"Issue {i}: {issue.severity.value.capitalize()}", level=3)
        
        # Add issue text
        doc.add_paragraph(f"Text: \"{issue.text}\"")
        
        # Add comment if available
        if issue.comment:
            comment_para = doc.add_paragraph("Comment: ")
            comment_para.add_run(issue.comment).italic = True
        
        # Add suggestion if available
        if issue.suggestion:
            suggestion_para = doc.add_paragraph("Suggestion: ")
            suggestion_run = suggestion_para.add_run(issue.suggestion)
            suggestion_run.bold = True
            suggestion_run.font.color.rgb = RGBColor(0, 0, 139)  # Dark blue
        
        doc.add_paragraph()  # Add space between issues
    
    # Save the document to a bytes buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer.getvalue()

def create_summary_markdown(
    issues: List[Dict[str, Any]],
    title: Optional[str] = None
) -> str:
    """
    Create a markdown summary of document issues.
    
    Args:
        issues: List of issues with keys:
               text, start, end, severity, comment, suggestion
        title: Optional document title
        
    Returns:
        str: Markdown summary
    """
    markdown = []
    
    # Add title
    if title:
        markdown.append(f"# {title}")
        markdown.append("")
    
    # Add summary section
    markdown.append("## Summary of Issues")
    markdown.append("")
    
    # Group issues by severity
    issues_by_severity = {
        "critical": [],
        "high": [],
        "medium": [],
        "low": []
    }
    
    for issue in issues:
        severity = issue.get("severity", "medium").lower()
        if severity not in issues_by_severity:
            severity = "medium"
        issues_by_severity[severity].append(issue)
    
    # Add summary table
    markdown.append("| Severity | Count | Description |")
    markdown.append("| --- | --- | --- |")
    
    for severity, severity_issues in issues_by_severity.items():
        count = len(severity_issues)
        if count > 0:
            description = {
                "critical": "Critical issues requiring immediate attention",
                "high": "Significant issues that should be addressed",
                "medium": "Moderate issues to consider",
                "low": "Minor issues for awareness"
            }.get(severity, "")
            
            markdown.append(f"| {severity.capitalize()} | {count} | {description} |")
    
    markdown.append("")
    
    # Add detailed issues section
    markdown.append("## Detailed Issues")
    markdown.append("")
    
    # Add issues by severity
    for severity in ["critical", "high", "medium", "low"]:
        severity_issues = issues_by_severity[severity]
        if severity_issues:
            markdown.append(f"### {severity.capitalize()} Issues")
            markdown.append("")
            
            for i, issue in enumerate(severity_issues, 1):
                markdown.append(f"#### Issue {i}")
                markdown.append(f"- **Text**: \"{issue['text']}\"")
                
                if issue.get("comment"):
                    markdown.append(f"- **Comment**: {issue['comment']}")
                
                if issue.get("suggestion"):
                    markdown.append(f"- **Suggestion**: {issue['suggestion']}")
                
                markdown.append("")
    
    return "\n".join(markdown)
