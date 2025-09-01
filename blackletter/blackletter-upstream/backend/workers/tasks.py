import asyncio
import logging
from datetime import datetime
from typing import Any, Dict, List

from ..db.session import AsyncSessionLocal
from ..jobs import crud
from ..models.schemas import AnalysisIssue, AnalysisResult, JobStatus
from ..services.storage import storage_service
from .celery_app import celery_app

logger = logging.getLogger(__name__)


@celery_app.task(bind=True, max_retries=3, default_retry_delay=60)
def process_contract(self, job_id: str):
    """
    Asynchronous Celery task to process a single contract analysis job.
    This task is designed to be idempotent and resilient to failures.
    """

    async def _run_task():
        """Inner async function to handle the actual processing."""
        start_time = datetime.utcnow()

        async with AsyncSessionLocal() as db:
            try:
                job = await crud.get_job_by_id(db, job_id)
                if not job:
                    logger.error(f"Job with ID {job_id} not found. Aborting task.")
                    return

                if job.status in [JobStatus.COMPLETED, JobStatus.FAILED]:
                    logger.warning(
                        f"Job {job_id} is already in terminal state ({job.status}). Skipping."
                    )
                    return

                await crud.update_job_status(
                    db,
                    job_id,
                    JobStatus.PROCESSING,
                    processing_step="Starting analysis",
                )

                logger.info(f"Job {job_id}: Downloading file {job.file_object_key}")
                file_content = await storage_service.get_file(job.file_object_key)

                await crud.update_job_status(
                    db, job_id, JobStatus.PROCESSING, processing_step="Extracting text"
                )

                extracted_text = await extract_text_from_file(
                    file_content, job.original_filename
                )

                await crud.update_job_status(
                    db,
                    job_id,
                    JobStatus.PROCESSING,
                    processing_step="Running compliance checks",
                )

                rule_issues = await run_gdpr_rule_engine(
                    extracted_text, job.contract_type, job.jurisdiction
                )

                await crud.update_job_status(
                    db, job_id, JobStatus.PROCESSING, processing_step="AI analysis"
                )

                llm_issues = await run_llm_analysis(extracted_text, rule_issues)

                await crud.update_job_status(
                    db,
                    job_id,
                    JobStatus.PROCESSING,
                    processing_step="Generating report",
                )

                all_issues = rule_issues + llm_issues
                compliant_issues = [i for i in all_issues if i.compliant]
                overall_score = (
                    (len(compliant_issues) / len(all_issues)) * 100
                    if all_issues
                    else 100
                )

                processing_time = (datetime.utcnow() - start_time).total_seconds()

                analysis_result = AnalysisResult(
                    summary=f"Analysis completed. {len(compliant_issues)}/{len(all_issues)} requirements met.",
                    overall_score=overall_score,
                    clauses_found=len(
                        [issue for issue in all_issues if issue.compliant]
                    ),
                    issues_detected=len(
                        [issue for issue in all_issues if not issue.compliant]
                    ),
                    issues=all_issues,
                    processing_time_seconds=processing_time,
                )

                report_file_key = await generate_pdf_report(job_id, analysis_result)

                await crud.update_job_result(
                    db=db,
                    job_id=job_id,
                    result=analysis_result,
                    report_file_key=report_file_key,
                    processing_time=processing_time,
                )

                logger.info(
                    f"Successfully processed job {job_id} in {processing_time:.2f} seconds."
                )

            except Exception as exc:
                logger.error(f"Error processing job {job_id}: {exc}", exc_info=True)

                await crud.update_job_status(
                    db=db,
                    job_id=job_id,
                    status=JobStatus.FAILED,
                    error_message=str(exc),
                )

                if self.request.retries < self.max_retries:
                    logger.info(
                        f"Retrying job {job_id} (attempt {self.request.retries + 1}/{self.max_retries})"
                    )
                    raise self.retry(exc=exc)

    asyncio.run(_run_task())


async def extract_text_from_file(file_content: bytes, filename: str) -> str:
    """Extract text content from uploaded file."""
    try:
        if filename.lower().endswith(".pdf"):
            return await extract_text_from_pdf(file_content)
        elif filename.lower().endswith((".docx", ".doc")):
            return await extract_text_from_docx(file_content)
        else:
            raise ValueError(f"Unsupported file type: {filename}")
    except Exception as e:
        logger.error(f"Failed to extract text from {filename}: {e}")
        raise


async def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text from PDF file using PyMuPDF."""
    import io

    import fitz

    try:
        pdf_document = fitz.open(stream=file_content, filetype="pdf")
        text = ""

        for page_num in range(pdf_document.page_count):
            page = pdf_document[page_num]
            page_text = page.get_text()
            if len(page_text.strip()) < 50:
                logger.info(
                    f"Page {page_num + 1} appears to be scanned, attempting OCR"
                )
                page_text = await ocr_pdf_page(page)
            text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"

        pdf_document.close()
        return text.strip()

    except Exception as e:
        logger.error(f"PDF extraction failed: {e}")
        raise ValueError(f"Failed to extract text from PDF: {e}")


async def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX file."""
    import io

    from docx import Document

    try:
        doc = Document(io.BytesIO(file_content))
        text = ""

        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"

        return text.strip()

    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        raise ValueError(f"Failed to extract text from DOCX: {e}")


async def ocr_pdf_page(page) -> str:
    """Perform OCR on a PDF page using Tesseract."""
    try:
        import io

        import pytesseract
        from PIL import Image

        pix = page.get_pixmap()
        img_data = pix.tobytes("png")
        image = Image.open(io.BytesIO(img_data))
        text = pytesseract.image_to_string(image, config="--psm 6")
        return text

    except ImportError:
        logger.warning("pytesseract not available, skipping OCR")
        return ""
    except Exception as e:
        logger.error(f"OCR failed: {e}")
        return ""


async def run_gdpr_rule_engine(
    text: str, contract_type, jurisdiction
) -> List[AnalysisIssue]:
    """Run rule-based GDPR compliance checks."""
    import re
    from pathlib import Path

    import yaml

    try:
        rules_file = Path(__file__).parent.parent / "rules" / "gdpr_playbook.yaml"

        if not rules_file.exists():
            rules = get_default_gdpr_rules()
        else:
            with open(rules_file, "r") as f:
                rules = yaml.safe_load(f)

        issues = []
        text_lower = text.lower()

        for rule in rules:
            rule_id = rule["id"]
            description = rule["description"]
            pattern = rule["pattern"]
            required = rule.get("required", True)
            severity = rule.get("severity", "medium")

            matches = re.findall(pattern, text, re.IGNORECASE | re.DOTALL)

            if required:
                if matches:
                    citation = (
                        matches[0]
                        if isinstance(matches[0], str)
                        else str(matches[0])[:200]
                    )
                    issues.append(
                        AnalysisIssue(
                            rule_id=rule_id,
                            description=description,
                            compliant=True,
                            severity=severity,
                            details=rule.get("if_found", "Required clause found."),
                            citation=f'"{citation}..."',
                            recommendation=None,
                        )
                    )
                else:
                    issues.append(
                        AnalysisIssue(
                            rule_id=rule_id,
                            description=description,
                            compliant=False,
                            severity=severity,
                            details=rule.get(
                                "if_missing", "Required clause not found."
                            ),
                            citation=None,
                            recommendation=rule.get(
                                "recommendation",
                                f"Add {description.lower()} clause to contract.",
                            ),
                        )
                    )
            else:
                if matches:
                    for match in matches[:3]:
                        issues.append(
                            AnalysisIssue(
                                rule_id=f"{rule_id}_{len(issues)}",
                                description=f"{description}: '{match}'",
                                compliant=False,
                                severity="low",
                                details=rule.get(
                                    "warning",
                                    f"Found potentially problematic term: {match}",
                                ),
                                citation=f'"{match}"',
                                recommendation=rule.get(
                                    "recommendation",
                                    f'Consider replacing "{match}" with more specific language.',
                                ),
                            )
                        )

        logger.info(f"Rule engine completed: {len(issues)} issues found")
        return issues

    except Exception as e:
        logger.error(f"Rule engine failed: {e}")
        return [
            AnalysisIssue(
                rule_id="rule_engine_error",
                description="Rule Engine Error",
                compliant=False,
                severity="high",
                details=f"Rule engine failed to complete: {str(e)}",
                citation=None,
                recommendation="Manual review required due to analysis error.",
            )
        ]


def get_default_gdpr_rules() -> List[dict]:
    """Return default GDPR compliance rules."""
    return [
        {
            "id": "breach_notification",
            "description": "Breach notification within 72 hours",
            "pattern": r"(breach|incident).*72\s*hour",
            "required": True,
            "severity": "high",
            "if_missing": "No clause requiring breach notification within 72 hours was found.",
            "if_found": "Includes a breach notification clause.",
            "recommendation": (
                'Add clause: "Processor shall notify Controller of any personal data breach within 72 hours of becoming aware of it."'
            ),
        },
        {
            "id": "data_deletion",
            "description": "Deletion or return of data on contract termination",
            "pattern": r"(delete|return).*(terminat|end|expir)",
            "required": True,
            "severity": "high",
            "if_missing": "No clause for deletion or return of personal data upon termination.",
            "if_found": "Includes end-of-contract data deletion/return clause.",
            "recommendation": (
                "Add clause requiring deletion or return of all personal data upon contract termination."
            ),
        },
        {
            "id": "subprocessor_consent",
            "description": "Sub-processor engagement requires prior consent",
            "pattern": r"sub.?process.*(consent|approval|written)",
            "required": True,
            "severity": "medium",
            "if_missing": "No clause requiring consent for sub-processor engagement.",
            "if_found": "Includes sub-processor consent requirement.",
            "recommendation": (
                "Add clause requiring prior written consent before engaging any sub-processors."
            ),
        },
        {
            "id": "security_measures",
            "description": "Appropriate technical and organizational security measures",
            "pattern": r"(security|technical|organizational).*(measure|safeguard)",
            "required": True,
            "severity": "high",
            "if_missing": "No clause describing security measures.",
            "if_found": "Includes security measures clause.",
            "recommendation": (
                "Add detailed clause describing technical and organizational security measures."
            ),
        },
        {
            "id": "data_subject_rights",
            "description": "Assistance with data subject rights requests",
            "pattern": r"(data subject|individual).*(right|request|access)",
            "required": True,
            "severity": "medium",
            "if_missing": "No clause for assisting with data subject rights.",
            "if_found": "Includes data subject rights assistance clause.",
            "recommendation": (
                "Add clause requiring processor to assist with data subject rights requests."
            ),
        },
        {
            "id": "vague_terms",
            "description": "Vague terms that should be more specific",
            "pattern": r"\b(reasonable|appropriate|adequate|suitable)\b",
            "required": False,
            "severity": "low",
            "warning": 'Uses vague term "{}" which could be made more specific.',
            "recommendation": "Replace vague terms with specific, measurable requirements.",
        },
    ]


async def run_llm_analysis(
    text: str, rule_issues: List[AnalysisIssue]
) -> List[AnalysisIssue]:
    """Run LLM-based analysis for complex compliance evaluation."""
    try:
        import openai

        from ..core.config import settings

        openai.api_key = settings.OPENAI_API_KEY
        llm_issues: List[AnalysisIssue] = []

        for issue in rule_issues:
            if not issue.compliant and issue.severity == "high":
                enhanced_issue = await analyze_missing_clause_with_llm(text, issue)
                if enhanced_issue:
                    llm_issues.append(enhanced_issue)
            elif issue.compliant and issue.citation:
                verified_issue = await verify_clause_adequacy_with_llm(
                    issue.citation, issue
                )
                if verified_issue:
                    llm_issues.append(verified_issue)

        logger.info(f"LLM analysis completed: {len(llm_issues)} additional insights")
        return llm_issues

    except Exception as e:
        logger.error(f"LLM analysis failed: {e}")
        return []


async def analyze_missing_clause_with_llm(
    text: str, issue: AnalysisIssue
) -> AnalysisIssue | None:
    """Use LLM to analyze missing clauses and provide detailed recommendations."""
    try:
        import openai

        prompt = f"""
    As a GDPR compliance expert, analyze this contract excerpt for the missing requirement:

    Missing Requirement: {issue.description}

    Contract Text (first 2000 chars):
    {text[:2000]}

    Provide:
    1. Confirmation that this requirement is indeed missing
    2. Specific impact of this missing clause
    3. Exact wording suggestion for the missing clause

    Format your response as:
    MISSING: Yes/No
    IMPACT: [Impact description]
    SUGGESTED_CLAUSE: [Exact wording]
    """

        response = await openai.ChatCompletion.acreate(
            model="gpt-3.5-turbo",
            messages=[
                {
                    "role": "system",
                    "content": "You are a GDPR compliance expert providing precise legal analysis.",
                },
                {"role": "user", "content": prompt},
            ],
            max_tokens=500,
            temperature=0.1,
        )

        analysis = response.choices[0].message.content

        if "MISSING: Yes" in analysis:
            impact_start = analysis.find("IMPACT:") + 7
            impact_end = analysis.find("SUGGESTED_CLAUSE:")
            impact = analysis[impact_start:impact_end].strip()

            clause_start = analysis.find("SUGGESTED_CLAUSE:") + 17
            suggested_clause = analysis[clause_start:].strip()

            return AnalysisIssue(
                rule_id=f"{issue.rule_id}_llm_enhanced",
                description=f"{issue.description} (LLM Analysis)",
                compliant=False,
                severity=issue.severity,
                details=f"LLM Analysis: {impact}",
                citation=None,
                recommendation=f"Suggested clause: {suggested_clause}",
            )

    except Exception as e:
        logger.error(f"LLM missing clause analysis failed: {e}")

    return None


async def verify_clause_adequacy_with_llm(
    citation: str, issue: AnalysisIssue
) -> AnalysisIssue | None:
    """Use LLM to verify if found clauses are actually adequate."""
    try:
        import openai

        prompt = f"""
    As a GDPR compliance expert, evaluate if this contract clause adequately meets the requirement:

    Requirement: {issue.description}
    Contract Clause: "{citation}"

    Analyze:
    1. Does this clause fully satisfy the GDPR requirement?
    2. Are there any gaps or weaknesses?
    3. What improvements could be made?

    Format your response as:
    ADEQUATE: Yes/No
    ANALYSIS: [Your analysis]
    IMPROVEMENTS: [Suggested improvements if any]
    """

        response = await openai.ChatCompletion.acreate(
            model="gpt-3.5-turbo",
            messages=[
                {
                    "role": "system",
                    "content": "You are a GDPR compliance expert providing detailed clause analysis.",
                },
                {"role": "user", "content": prompt},
            ],
            max_tokens=400,
            temperature=0.1,
        )

        analysis = response.choices[0].message.content

        if "ADEQUATE: No" in analysis:
            analysis_start = analysis.find("ANALYSIS:") + 9
            analysis_end = analysis.find("IMPROVEMENTS:")
            llm_analysis = analysis[analysis_start:analysis_end].strip()

            improvements_start = analysis.find("IMPROVEMENTS:") + 13
            improvements = analysis[improvements_start:].strip()

            return AnalysisIssue(
                rule_id=f"{issue.rule_id}_adequacy_check",
                description=f"{issue.description} - Adequacy Concern",
                compliant=False,
                severity="medium",
                details=f"LLM found clause inadequate: {llm_analysis}",
                citation=citation,
                recommendation=improvements,
            )

    except Exception as e:
        logger.error(f"LLM adequacy check failed: {e}")

    return None


async def generate_pdf_report(job_id: str, analysis_result: AnalysisResult) -> str:
    """Generate PDF report and return S3 object key."""
    try:
        from datetime import datetime

        import weasyprint
        from jinja2 import Template

        html_template = Template(
            """
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="utf-8">
        <title>GDPR Compliance Report</title>
        <style>
            body { font-family: Arial, sans-serif; margin: 40px; }
            .header { text-align: center; border-bottom: 2px solid #333; padding-bottom: 20px; }
            .summary { background-color: #f5f5f5; padding: 20px; margin: 20px 0; border-radius: 5px; }
            .score { font-size: 24px; font-weight: bold; color: #2c5aa0; }
            .issue { margin: 20px 0; padding: 15px; border-left: 4px solid #ddd; }
            .compliant { border-left-color: #28a745; }
            .non-compliant { border-left-color: #dc3545; }
            .citation { background-color: #f8f9fa; padding: 10px; font-style: italic; margin: 10px 0; }
            .recommendation { background-color: #fff3cd; padding: 10px; margin: 10px 0; border-radius: 3px; }
            .footer { margin-top: 40px; padding-top: 20px; border-top: 1px solid #ddd; font-size: 12px; color: #666; }
        </style>
    </head>
    <body>
        <div class="header">
            <h1>GDPR Compliance Report</h1>
            <p>Generated on {{ date }}</p>
            <p>Job ID: {{ job_id }}</p>
        </div>

        <div class="summary">
            <h2>Executive Summary</h2>
            <div class="score">Overall Compliance Score: {{ "%.1f"|format(analysis_result.overall_score) }}%</div>
            <p>{{ analysis_result.summary }}</p>
            <p><strong>Clauses Found:</strong> {{ analysis_result.clauses_found }}</p>
            <p><strong>Issues Detected:</strong> {{ analysis_result.issues_detected }}</p>
            {% if analysis_result.processing_time_seconds %}
            <p><strong>Processing Time:</strong> {{ "%.2f"|format(analysis_result.processing_time_seconds) }} seconds</p>
            {% endif %}
        </div>

        <h2>Detailed Analysis</h2>
        {% for issue in analysis_result.issues %}
        <div class="issue {{ 'compliant' if issue.compliant else 'non-compliant' }}">
            <h3>
                {% if issue.compliant %}✅{% else %}❌{% endif %}
                {{ issue.description }}
            </h3>
            <p><strong>Status:</strong> {{ "Compliant" if issue.compliant else "Non-Compliant" }}</p>
            <p><strong>Severity:</strong> {{ issue.severity.title() }}</p>
            <p>{{ issue.details }}</p>

            {% if issue.citation %}
            <div class="citation">
                <strong>Contract Reference:</strong><br>
                {{ issue.citation }}
            </div>
            {% endif %}

            {% if issue.recommendation %}
            <div class="recommendation">
                <strong>Recommendation:</strong><br>
                {{ issue.recommendation }}
            </div>
            {% endif %}
        </div>
        {% endfor %}

        <div class="footer">
            <p>This report was generated by Blackletter GDPR Contract Checker.</p>
            <p>Please consult with legal counsel for definitive compliance advice.</p>
            <p>Report data will be automatically deleted after 30 days unless requested otherwise.</p>
        </div>
    </body>
    </html>
            """
        )

        html_content = html_template.render(
            job_id=job_id,
            date=datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC"),
            analysis_result=analysis_result,
        )

        pdf_bytes = weasyprint.HTML(string=html_content).write_pdf()

        report_key = f"reports/{job_id}/compliance_report.pdf"

        storage_service.s3_client.put_object(
            Bucket=storage_service.bucket_name,
            Key=report_key,
            Body=pdf_bytes,
            ContentType="application/pdf",
            ServerSideEncryption="AES256",
        )

        logger.info(f"Generated PDF report for job {job_id}: {report_key}")
        return report_key

    except Exception as e:
        logger.error(f"Failed to generate PDF report for job {job_id}: {e}")
        return None
